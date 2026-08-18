"""各类文档的文本抽取。

每种格式一个函数，统一返回 (text: str, meta: dict)。
任何格式的单点失败都应抛出 ExtractError，由上层决定降级为 status='error'。
"""
import csv
import io
import logging
import os

logger = logging.getLogger("documents.extractors")


class ExtractError(Exception):
    """抽取失败（文件损坏、格式不支持、依赖缺等）。"""


# ----------------------------------------------------------------------------
# PDF
# ----------------------------------------------------------------------------
def extract_pdf(path: str):
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover
        raise ExtractError("缺少 pypdf 依赖") from e
    try:
        reader = PdfReader(path)
    except Exception as e:
        raise ExtractError(f"PDF 解析失败：{e}") from e
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception as e:
            logger.warning("PDF 第 %d 页抽取异常：%s", i + 1, e)
            txt = ""
        if txt:
            parts.append(txt)
    text = "\n\n".join(parts).strip()
    return text, {"page_count": len(reader.pages)}


# ----------------------------------------------------------------------------
# Word (.docx)
# ----------------------------------------------------------------------------
def extract_docx(path: str):
    try:
        import docx
    except ImportError as e:  # pragma: no cover
        raise ExtractError("缺少 python-docx 依赖") from e
    try:
        doc = docx.Document(path)
    except Exception as e:
        raise ExtractError(f"Word 解析失败：{e}") from e

    blocks = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        # 标题作为上下文保留进正文流：分块时自然带上"章节"信息，
        # 同时作为该段内容的元数据线索（Heading1/2/3 -> #/##/###）。
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading") or style.startswith("标题"):
            level = 1
            m = __import__("re").search(r"(\d)", style)
            if m:
                level = min(int(m.group(1)), 3)
            blocks.append(("#" * level) + " " + text)
        else:
            blocks.append(text)

    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                blocks.append(line)
    return "\n".join(blocks).strip(), {
        "paragraph_count": len(doc.paragraphs),
        "table_count": table_count,
    }


# ----------------------------------------------------------------------------
# Excel (.xlsx / .xls)
# ----------------------------------------------------------------------------
def extract_xlsx(path: str):
    if path.lower().endswith(".xls"):
        raise ExtractError("不支持旧版 .xls，请另存为 .xlsx 后上传")
    try:
        import openpyxl
    except ImportError as e:  # pragma: no cover
        raise ExtractError("缺少 openpyxl 依赖") from e
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as e:
        raise ExtractError(f"Excel 解析失败：{e}") from e

    sheets = []
    for ws in wb.worksheets:
        rows = []
        for r in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).strip() for v in r]
            line = "\t".join(cells).strip()
            if line:
                rows.append(line)
        if rows:
            sheets.append(f"# 工作表：{ws.title}\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets).strip(), {"sheets": [ws.title for ws in wb.worksheets]}


# ----------------------------------------------------------------------------
# CSV
# ----------------------------------------------------------------------------
def extract_csv(path: str):
    try:
        with open(path, "r", encoding="utf-8-sig", newline="") as f:
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
            except Exception:
                dialect = csv.excel
            reader = csv.reader(f, dialect)
            rows = []
            for r in reader:
                line = "\t".join(c.strip() for c in r)
                if line:
                    rows.append(line)
    except UnicodeDecodeError:
        with open(path, "r", encoding="gbk", newline="") as f:
            reader = csv.reader(f)
            rows = ["\t".join(c.strip() for c in r) for r in reader]
    except Exception as e:
        raise ExtractError(f"CSV 解析失败：{e}") from e
    return "\n".join(rows).strip(), {"row_count": len(rows)}


# ----------------------------------------------------------------------------
# TXT / MD
# ----------------------------------------------------------------------------
def extract_text(path: str):
    last_err = None
    for enc in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read().strip(), {"encoding": enc}
        except (UnicodeDecodeError, UnicodeError) as e:
            last_err = e
            continue
    raise ExtractError(f"文本解码失败：{last_err}") from last_err


# ----------------------------------------------------------------------------
# 分发
# ----------------------------------------------------------------------------
_EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".xls": extract_xlsx,
    ".csv": extract_csv,
    ".txt": extract_text,
    ".md": extract_text,
}


def extract(path: str, file_type: str):
    """按 file_type 分发抽取，返回 (text, meta)。file_type 形如 'pdf'。"""
    ext = "." + file_type.lower().lstrip(".")
    fn = _EXTRACTORS.get(ext)
    if fn is None:
        raise ExtractError(f"不支持的文件类型：{file_type}")
    return fn(path)


def detect_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    return ext.lstrip(".")
