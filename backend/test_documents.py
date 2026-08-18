"""文档处理模块端到端测试。

生成 pdf/docx/xlsx/csv/txt 样例 -> 经 Flask test client 走完整 API：
upload -> 断言抽取/切片 -> list/detail/chunks -> delete。
同时验证不支持类型(.xls)与不支持扩展名的优雅降级。
所有测试数据最后清理（含磁盘文件）。
"""
import csv
import io
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app import create_app

# ---------------------------------------------------------------------------
# 样例文件构造
# ---------------------------------------------------------------------------
def make_pdf(path, text):
    esc = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = ("BT /F1 12 Tf 72 720 Td (%s) Tj ET" % esc).encode("latin-1")
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(out.tell())
        out.write(b"%d 0 obj\n" % i)
        out.write(body)
        out.write(b"\nendobj\n")
    xref_pos = out.tell()
    out.write(b"xref\n0 %d\n" % (len(objs) + 1))
    out.write(b"0000000000 65535 f \n")
    for off in offsets:
        out.write(("%010d 00000 n \n" % off).encode("latin-1"))
    out.write(b"trailer\n<< /Size %d /Root 1 0 R >>\n" % (len(objs) + 1))
    out.write(b"startxref\n%d\n%%%%EOF" % xref_pos)
    with open(path, "wb") as f:
        f.write(out.getvalue())


def make_docx(path):
    import docx
    d = docx.Document()
    d.add_paragraph("第一章 项目背景")
    d.add_paragraph("本项目用于验证文档解析能力，支持多种格式。")
    d.add_paragraph("第二章 技术方案")
    d.add_paragraph("采用前后端分离架构，后端使用 Flask。")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "模块"
    t.cell(0, 1).text = "说明"
    t.cell(1, 0).text = "documents"
    t.cell(1, 1).text = "文档处理"
    d.save(path)


def make_xlsx(path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "销售"
    ws.append(["月份", "销售额"])
    ws.append(["一月", 100])
    ws.append(["二月", 200])
    wb.save(path)


def make_csv(path):
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["name", "age"])
        w.writerow(["Alice", 30])
        w.writerow(["Bob", 25])


def make_txt(path):
    with open(path, "w", encoding="utf-8") as f:
        f.write("这是一段纯文本。\n\n用于测试文本抽取与切片功能。\n\n分段内容会被正确切分。")


def upload(client, path, filename):
    with open(path, "rb") as f:
        data = {"file": (f, filename)}
        return client.post("/api/documents/upload", data=data, content_type="multipart/form-data")


PASS = 0
FAIL = 0


def check(name, cond, extra=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  [PASS] {name} {extra}")
    else:
        FAIL += 1
        print(f"  [FAIL] {name} {extra}")


def main():
    app = create_app()
    # 单测用独立上传目录，避免污染真实数据
    tmp = tempfile.mkdtemp(prefix="doc_test_")
    app.config["DOC_UPLOAD_DIR"] = tmp
    import config
    config.Config.DOC_UPLOAD_DIR = tmp

    created_ids = []
    with app.test_client() as c:
        # 列表初始为空（此测试库可能是真实库，只断言返回的是 list）
        r = c.get("/api/documents")
        check("列表接口返回 list", r.status_code == 200 and isinstance(r.get_json()["data"], list))

        # 各格式
        cases = [
            ("sample.pdf", "PDF 抽取", "PDF extraction sample text for testing."),
            ("sample.docx", "Word 抽取", None),
            ("sample.xlsx", "Excel 抽取", None),
            ("sample.csv", "CSV 抽取", None),
            ("sample.txt", "TXT 抽取", None),
        ]
        builders = {
            "sample.pdf": lambda p: make_pdf(p, "PDF extraction sample text for testing."),
            "sample.docx": make_docx,
            "sample.xlsx": make_xlsx,
            "sample.csv": make_csv,
            "sample.txt": make_txt,
        }
        for fname, label, _ in cases:
            p = os.path.join(tmp, fname)
            builders[fname](p)
            r = upload(c, p, fname)
            j = r.get_json()
            data = j["data"]
            created_ids.append(data["id"])
            check(f"{label}: http 200", r.status_code == 200, f"code={j['code']}")
            check(f"{label}: status=done", data["status"] == "done", f"status={data['status']} err={data['error']}")
            check(f"{label}: chunk_count>0", data["chunk_count"] > 0, f"chunks={data['chunk_count']}")

        # 详情含全文
        doc_id = created_ids[0]
        r = c.get(f"/api/documents/{doc_id}")
        d = r.get_json()["data"]
        check("详情含 text 字段", isinstance(d.get("text"), str) and len(d["text"]) > 0)
        check("详情含 meta_json", d.get("meta_json"))

        # chunks
        r = c.get(f"/api/documents/{doc_id}/chunks")
        chunks = r.get_json()["data"]
        check("chunks 接口返回列表", isinstance(chunks, list) and len(chunks) > 0)
        check("chunk 含 idx/content", all("idx" in x and "content" in x for x in chunks))

        # 不支持扩展名降级
        p = os.path.join(tmp, "bad.xyz")
        with open(p, "wb") as f:
            f.write(b"x")
        r = upload(c, p, "bad.xyz")
        check("不支持扩展名 -> 400", r.status_code == 400, f"code={r.get_json()['code']}")

        # .xls 不支持（缺依赖）
        p = os.path.join(tmp, "old.xls")
        with open(p, "wb") as f:
            f.write(b"\xd0\xcf\x11\xe0")  # OLE 头
        r = upload(c, p, "old.xls")
        j = r.get_json()
        check(".xls -> status=error 不崩溃", j["data"]["status"] == "error", f"err={j['data']['error']}")

        # 删除（含磁盘文件）
        r = c.delete(f"/api/documents/{doc_id}")
        check("删除接口 200", r.status_code == 200 and r.get_json()["data"]["deleted"] is True)
        # 磁盘文件应被清理
        left = [f for f in os.listdir(tmp) if f.startswith(doc_id)]
        check("删除后磁盘文件清理", len(left) == 0, f"残留={left}")

    # 清理其余测试文档记录
    with app.test_client() as c:
        for did in created_ids[1:]:
            c.delete(f"/api/documents/{did}")

    print(f"\n结果：PASS={PASS}  FAIL={FAIL}")
    sys.exit(1 if FAIL else 0)


if __name__ == "__main__":
    main()
